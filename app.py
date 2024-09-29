from flask import Flask, render_template, request, jsonify, send_file, abort
import os
import docx
from reportlab.pdfgen import canvas
import json
import tempfile
from urllib.parse import quote, unquote
import random
import csv

app = Flask(__name__)

# Load the questions from the CSV file
questions_db = {}

# Updated to load questions correctly for all subjects
with open(r'C:\Users\risha\PycharmProjects\Cufront\quizzy\questions.csv', 'r') as csvfile:
    reader = csv.DictReader(csvfile)
    for row in reader:
        subject = row['Subject']
        question = row['Question']
        options = [row['Option 1'], row['Option 2'], row['Option 3'], row['Option 4']]
        correct_answer = row['Correct Answer']
        if subject not in questions_db:
            questions_db[subject] = []
        questions_db[subject].append({
            'question': question,
            'options': options,
            'correct_answer': correct_answer
        })

@app.route('/')
def index():
    return render_template("index.html", questions_db=questions_db)

@app.route('/generate_quiz', methods=['POST'])
def generate_quiz():
    subject = request.form.get('subject')
    num_questions = int(request.form.get('num_questions'))
    include_answers = request.form.get('include_answers') == 'on'  # Check if correct answers should be included

    if num_questions > len(questions_db[subject]):
        return abort(400, description="Not enough questions available")

    selected_questions = random.sample(questions_db[subject], num_questions)  # Randomly select questions

    # URL-encode the questions and include_answers
    encoded_questions = quote(json.dumps(selected_questions))

    return render_template('results.html', questions=selected_questions, subject=subject,
                           encoded_questions=encoded_questions, include_answers=include_answers)

@app.route('/download_quiz/word', methods=['GET'])
def download_quiz_word():
    subject = request.args.get('subject')
    encoded_questions = request.args.get('questions')
    include_answers = request.args.get('include_answers', 'false') == 'True'
    questions = json.loads(unquote(encoded_questions))

    doc = docx.Document()
    doc.add_heading(f'{subject} Quiz', level=1)

    for question in questions:
        doc.add_paragraph(question['question'])
        for i, option in enumerate(question['options']):
            doc.add_paragraph(f'{chr(65 + i)}. {option}')
        if include_answers:
            doc.add_paragraph(f'Correct Answer: {question["correct_answer"]}')

    temp_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    doc.save(temp_file.name)
    temp_file.close()

    return send_file(temp_file.name, as_attachment=True, download_name=f'{subject}_quiz.docx')


@app.route('/download_quiz/pdf', methods=['GET'])
def download_quiz_pdf():
    subject = request.args.get('subject')
    encoded_questions = request.args.get('questions')
    include_answers = request.args.get('include_answers', 'false') == 'True'
    questions = json.loads(unquote(encoded_questions))

    pdf_file = f'{subject}_quiz.pdf'
    c = canvas.Canvas(pdf_file)
    c.setFont("Helvetica", 12)
    c.drawString(100, 750, f'{subject} Quiz')
    y = 720

    # Initialize question counter
    question_count = 0

    for index, question in enumerate(questions):
        # Check if we need to create a new page after every 5 questions
        if question_count >= 5:
            c.showPage()
            c.setFont("Helvetica", 12)
            c.drawString(100, 750, f'{subject} Quiz')  # Heading for new page
            y = 700  # Set Y position lower to create a gap below the heading
            question_count = 0  # Reset the question count

        # Draw the numbered question
        c.drawString(100, y, f'{index + 1}. {question["question"]}')  # Numbering the questions
        y -= 20
        for i, option in enumerate(question['options']):
            # Draw each option
            c.drawString(100, y, f'{chr(65 + i)}. {option}')
            y -= 20

        if include_answers:
            c.drawString(100, y, f'Correct Answer: {question["correct_answer"]}')
            y -= 20

        # Add some extra space between questions
        y -= 10
        question_count += 1  # Increment question count

    c.save()

    return send_file(pdf_file, as_attachment=True, download_name=f'{subject}_quiz.pdf')


@app.route('/download_quiz/json', methods=['GET'])
def download_quiz_json():
    subject = request.args.get('subject')
    encoded_questions = request.args.get('questions')
    questions = json.loads(unquote(encoded_questions))

    return jsonify(questions)

if __name__ == '__main__':
    app.run(debug=True)
